"""Microsoft Word document reader (.docx).

Optional dependency: ``python-docx>=1.1``.  Install with:

    pip install "lean-ai[reference]"

or manually:

    pip install python-docx

Paragraphs are grouped by their Word style headings (Heading 1, Heading 2,
etc.) so each logical section becomes a separate set of chunks.  Tables
are rendered as markdown tables so tabular data (e.g. reference tables)
survives the conversion in a readable form.

The module-level ``docx_to_markdown`` function is reused outside the
reference library — ``read_file`` calls it to support ``.docx`` inputs
and the ``/api/workspace/convert-docx`` endpoint uses it to produce a
deterministic markdown copy of a Word document.
"""

import logging
from pathlib import Path

from lean_ai.reference.chunker import chunk_prose_configured
from lean_ai.reference.readers.base import DocumentReader, ReferenceChunk

logger = logging.getLogger(__name__)

# Markdown cell placeholder — a regular space renders as an empty cell but
# some markdown parsers collapse empty cells and misalign columns.  A
# non-breaking space keeps the column count stable without being visible.
_EMPTY_CELL = "\u00a0"


def _heading_level(style_name: str) -> int | None:
    """Return the heading depth for a Word paragraph style (1-6), else None."""
    if not style_name.startswith("Heading "):
        return None
    suffix = style_name[len("Heading "):].strip()
    if not suffix.isdigit():
        return None
    level = int(suffix)
    return level if 1 <= level <= 6 else None


def _render_markdown_table(rows: list[list[str]]) -> str:
    """Render a rectangular list of string rows as a markdown table.

    The first row is used as the header.  Short rows are padded with
    ``_EMPTY_CELL`` so every row has the same column count.  Pipe
    characters inside cells are escaped so the table parses correctly.
    """
    if not rows:
        return ""

    width = max(len(r) for r in rows)

    def _normalise(row: list[str]) -> list[str]:
        padded = list(row) + [_EMPTY_CELL] * (width - len(row))
        return [
            (cell.replace("|", "\\|").replace("\n", " ").strip() or _EMPTY_CELL)
            for cell in padded
        ]

    header = _normalise(rows[0])
    body = [_normalise(r) for r in rows[1:]]

    lines = ["| " + " | ".join(header) + " |"]
    lines.append("| " + " | ".join(["---"] * width) + " |")
    for row in body:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def docx_to_markdown(path: Path) -> str:
    """Convert a ``.docx`` file to a markdown string.

    Renders Word headings (1-6) as ``#``/``##``/...  headers, regular
    paragraphs as plain text, and tables as proper markdown tables.

    Raises:
        ImportError: if ``python-docx`` is not installed.  The caller is
            responsible for catching this and surfacing an install hint
            in a user-appropriate way.
        Exception: any parse error raised by ``python-docx`` propagates
            up (e.g. password-protected documents raise a generic error).
    """
    from docx import Document

    doc = Document(str(path))
    blocks: list[str] = []

    # Body: paragraphs with heading-style awareness
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        level = _heading_level(style_name)
        if level is not None:
            blocks.append(f"{'#' * level} {text}")
        else:
            blocks.append(text)

    # Tables render as markdown tables, appended after the body so they
    # don't interrupt narrative flow.  This mirrors the original reader's
    # behaviour of collecting tables separately.
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        # Drop trailing all-empty rows for cleaner output.
        while rows and not any(cell for cell in rows[-1]):
            rows.pop()
        if not rows:
            continue
        table_md = _render_markdown_table(rows)
        if table_md:
            blocks.append(table_md)

    return "\n\n".join(blocks)


class DocxReader(DocumentReader):
    """Reader for ``.docx`` files (requires ``python-docx``)."""

    @property
    def extensions(self) -> list[str]:
        return [".docx"]

    def read(self, path: Path, rel_path: str) -> list[ReferenceChunk]:
        try:
            markdown = docx_to_markdown(path)
        except ImportError as e:
            logger.warning(
                "Cannot read Word document %s — missing optional dependency: %s. "
                "Install with: pip install python-docx",
                path, e,
            )
            return []
        except Exception as e:
            logger.warning("Failed to open Word document %s: %s", path, e)
            return []

        doc_title = path.stem.replace("_", " ").replace("-", " ")

        # Split the flat markdown back into sections keyed by the nearest
        # preceding ``#``/``##`` heading, matching the previous reader's
        # chunk grouping so retrieval stays section-aware.  The first
        # Heading 1 becomes the doc title if the filename is generic.
        sections: list[tuple[str, list[str]]] = [("", [])]
        for block in markdown.split("\n\n"):
            if not block.strip():
                continue
            if block.startswith("#"):
                heading_text = block.lstrip("#").strip()
                # If the filename-derived title is still the default and
                # we see a Heading 1, promote it.
                default_title = path.stem.replace("_", " ").replace("-", " ")
                if block.startswith("# ") and doc_title == default_title:
                    doc_title = heading_text
                sections.append((heading_text, []))
            else:
                sections[-1][1].append(block)

        chunks: list[ReferenceChunk] = []
        chunk_index = 0
        for section_heading, body_blocks in sections:
            if not body_blocks:
                continue
            section_text = "\n\n".join(body_blocks)
            raw_chunks = chunk_prose_configured(section_text)
            for chunk_text in raw_chunks:
                if not chunk_text.strip():
                    continue
                chunks.append(ReferenceChunk(
                    doc_path=rel_path,
                    doc_title=doc_title,
                    section=section_heading,
                    content=chunk_text,
                    chunk_index=chunk_index,
                    format="docx",
                ))
                chunk_index += 1

        return chunks
