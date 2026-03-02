"""Microsoft Word document reader (.docx).

Optional dependency: ``python-docx>=1.1``.  Install with:

    pip install "lean-ai[knowledge]"

or manually:

    pip install python-docx

Paragraphs are grouped by their Word style headings (Heading 1, Heading 2,
etc.) so each logical section becomes a separate set of chunks.  Tables
are extracted cell-by-cell so tabular data (e.g. reference tables) is
preserved in a readable form.
"""

import logging
from pathlib import Path

from lean_ai.knowledge.chunker import chunk_prose
from lean_ai.knowledge.readers.base import DocumentReader, KnowledgeChunk

logger = logging.getLogger(__name__)


class DocxReader(DocumentReader):
    """Reader for ``.docx`` files (requires ``python-docx``)."""

    @property
    def extensions(self) -> list[str]:
        return [".docx"]

    def read(self, path: Path, rel_path: str) -> list[KnowledgeChunk]:
        try:
            from docx import Document
        except ImportError as e:
            logger.warning(
                "Cannot read Word document %s — missing optional dependency: %s. "
                "Install with: pip install python-docx",
                path, e,
            )
            return []

        try:
            doc = Document(str(path))
        except Exception as e:
            logger.warning("Failed to open Word document %s: %s", path, e)
            return []

        doc_title = path.stem.replace("_", " ").replace("-", " ")

        # Split paragraphs into sections based on heading styles.
        sections: list[tuple[str, list[str]]] = []
        current_heading = ""
        current_body: list[str] = []

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()

            if not text:
                continue

            if style_name.startswith("Heading"):
                # Flush the current section.
                if current_body:
                    sections.append((current_heading, current_body))
                    current_body = []

                # Use the first Heading 1 as the document title (if better
                # than the filename).
                default_title = path.stem.replace("_", " ").replace("-", " ")
                if style_name == "Heading 1" and doc_title == default_title:
                    doc_title = text

                current_heading = text
            else:
                current_body.append(text)

        # Flush the last section.
        if current_body:
            sections.append((current_heading, current_body))

        # Also extract table content.
        table_texts: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_texts.append(" | ".join(cells))

        if table_texts:
            sections.append(("Tables", table_texts))

        chunks: list[KnowledgeChunk] = []
        chunk_index = 0

        for section_heading, body_lines in sections:
            section_text = "\n\n".join(body_lines)
            raw_chunks = chunk_prose(section_text)
            for chunk_text in raw_chunks:
                if not chunk_text.strip():
                    continue
                chunks.append(KnowledgeChunk(
                    doc_path=rel_path,
                    doc_title=doc_title,
                    section=section_heading,
                    content=chunk_text,
                    chunk_index=chunk_index,
                    format="docx",
                ))
                chunk_index += 1

        return chunks
