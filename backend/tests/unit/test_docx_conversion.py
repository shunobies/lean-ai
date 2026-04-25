"""Tests for ``.docx`` → markdown conversion.

Covers the shared ``docx_to_markdown`` utility, ``read_file``'s new
``.docx`` branch, and the ``/api/workspace/convert-docx`` REST endpoint.

Fixtures are built in-memory with ``python-docx`` — no committed binary
files needed.  If ``python-docx`` is missing the tests are skipped.
"""

from __future__ import annotations

from pathlib import Path

import pytest

try:
    from docx import Document

    HAS_DOCX = True
except ImportError:  # pragma: no cover — environmental skip
    HAS_DOCX = False

pytestmark = pytest.mark.skipif(
    not HAS_DOCX,
    reason="python-docx not installed",
)


def _build_sample_docx(path: Path) -> None:
    """Create a small ``.docx`` fixture with headings, paragraphs, and a table."""
    doc = Document()
    doc.add_heading("Jane Q Candidate", level=1)
    doc.add_paragraph("Senior Software Engineer — based in Portland, OR.")
    doc.add_heading("Experience", level=2)
    doc.add_paragraph("Acme Corp, 2021 - present — led backend migration to Django.")
    doc.add_paragraph("BigCo, 2017 - 2021 — Python services, on-call rotation.")
    doc.add_heading("Skills", level=2)
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Language"
    table.rows[0].cells[1].text = "Years"
    table.rows[1].cells[0].text = "Python"
    table.rows[1].cells[1].text = "9"
    table.rows[2].cells[0].text = "Go"
    table.rows[2].cells[1].text = "3"
    doc.save(str(path))


# ── docx_to_markdown ─────────────────────────────────────────────────


def test_docx_to_markdown_renders_headings_and_paragraphs(tmp_path: Path) -> None:
    from lean_ai.reference.readers.docx import docx_to_markdown

    src = tmp_path / "resume.docx"
    _build_sample_docx(src)

    markdown = docx_to_markdown(src)

    assert "# Jane Q Candidate" in markdown
    assert "## Experience" in markdown
    assert "Senior Software Engineer" in markdown
    assert "Acme Corp, 2021 - present" in markdown


def test_docx_to_markdown_renders_proper_markdown_table(tmp_path: Path) -> None:
    from lean_ai.reference.readers.docx import docx_to_markdown

    src = tmp_path / "resume.docx"
    _build_sample_docx(src)

    markdown = docx_to_markdown(src)

    # Header row + separator + two body rows.
    assert "| Language | Years |" in markdown
    assert "| --- | --- |" in markdown
    assert "| Python | 9 |" in markdown
    assert "| Go | 3 |" in markdown


def test_docx_to_markdown_raises_import_error_when_docx_missing(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Caller is responsible for catching ImportError — verify it propagates."""
    import builtins

    from lean_ai.reference.readers import docx as docx_module

    src = tmp_path / "resume.docx"
    _build_sample_docx(src)

    real_import = builtins.__import__

    def _blocked_import(name, *args, **kwargs):
        if name == "docx" or name.startswith("docx."):
            raise ImportError("simulated: python-docx missing")
        return real_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", _blocked_import)

    with pytest.raises(ImportError):
        docx_module.docx_to_markdown(src)


# ── DocxReader reference-library integration ────────────────────────


def test_docx_reader_preserves_section_chunks(tmp_path: Path) -> None:
    """Reference indexing path still produces section-aware ReferenceChunks."""
    from lean_ai.reference.readers.docx import DocxReader

    src = tmp_path / "resume.docx"
    _build_sample_docx(src)

    chunks = DocxReader().read(src, "resume.docx")

    assert chunks, "DocxReader returned no chunks"
    sections = {c.section for c in chunks}
    # Heading 2s become section keys.
    assert "Experience" in sections
    assert "Skills" in sections
    # Chunk metadata looks sane.
    for c in chunks:
        assert c.doc_path == "resume.docx"
        assert c.format == "docx"
        assert c.chunk_index >= 0
        assert c.content.strip()


def test_docx_reader_promotes_heading_1_to_doc_title(tmp_path: Path) -> None:
    from lean_ai.reference.readers.docx import DocxReader

    src = tmp_path / "generic_filename.docx"
    _build_sample_docx(src)

    chunks = DocxReader().read(src, "generic_filename.docx")

    assert chunks
    assert chunks[0].doc_title == "Jane Q Candidate"


# ── read_file integration ────────────────────────────────────────────


@pytest.mark.asyncio
async def test_read_file_supports_docx(tmp_path: Path) -> None:
    from lean_ai.tools.file_ops import read_file

    src = tmp_path / "resume.docx"
    _build_sample_docx(src)

    result = await read_file("resume.docx", str(tmp_path))

    assert result.success
    assert result.output is not None
    # Output is line-numbered; strip to compare.
    body = "\n".join(line.split(" | ", 1)[-1] for line in result.output.splitlines())
    assert "Jane Q Candidate" in body
    assert "Experience" in body
    assert "Python" in body


@pytest.mark.asyncio
async def test_read_file_reports_parse_failure_for_corrupt_docx(tmp_path: Path) -> None:
    """A non-docx file with a .docx extension surfaces a clear parse error."""
    from lean_ai.tools.file_ops import read_file

    src = tmp_path / "fake.docx"
    src.write_text("this is not a real docx")

    result = await read_file("fake.docx", str(tmp_path))

    assert not result.success
    assert result.error is not None
    assert "Failed to parse .docx" in result.error


@pytest.mark.asyncio
async def test_read_file_preserves_binary_error_for_non_docx(tmp_path: Path) -> None:
    """Regression: .png still gets the binary-file error, not the .docx path."""
    from lean_ai.tools.file_ops import read_file

    src = tmp_path / "image.png"
    src.write_bytes(b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR")

    result = await read_file("image.png", str(tmp_path))

    assert not result.success
    assert result.error is not None
    assert "binary" in result.error.lower()


# ── REST endpoint ────────────────────────────────────────────────────


@pytest.fixture
def client():
    from fastapi.testclient import TestClient

    from lean_ai.main import app

    return TestClient(app)


def test_convert_docx_endpoint_success(tmp_path: Path, client) -> None:
    src = tmp_path / "master_resume.docx"
    _build_sample_docx(src)

    resp = client.post(
        "/api/workspace/convert-docx",
        json={
            "repo_root": str(tmp_path),
            "source_path": str(src),
            "output_filename": "acme_corp_senior_engineer_resume.md",
        },
    )

    assert resp.status_code == 200
    data = resp.json()
    assert data["success"] is True
    out = Path(data["output_path"])
    assert out.exists()
    content = out.read_text(encoding="utf-8")
    assert "# Jane Q Candidate" in content
    assert "| Python | 9 |" in content
    assert data["content_length"] == len(content)
    assert data["line_count"] >= 1


def test_convert_docx_endpoint_rejects_duplicate(tmp_path: Path, client) -> None:
    src = tmp_path / "master_resume.docx"
    _build_sample_docx(src)

    body = {
        "repo_root": str(tmp_path),
        "source_path": str(src),
        "output_filename": "acme_resume.md",
    }

    first = client.post("/api/workspace/convert-docx", json=body)
    assert first.status_code == 200

    # Second call without overwrite should 409.
    second = client.post("/api/workspace/convert-docx", json=body)
    assert second.status_code == 409
    detail = second.json()["detail"]
    assert detail["error"] == "exists"
    assert "acme_resume.md" in detail["output_path"]

    # Third call with overwrite=true succeeds.
    third = client.post(
        "/api/workspace/convert-docx",
        json={**body, "overwrite": True},
    )
    assert third.status_code == 200


def test_convert_docx_endpoint_rejects_non_docx_source(tmp_path: Path, client) -> None:
    src = tmp_path / "notes.txt"
    src.write_text("hello")

    resp = client.post(
        "/api/workspace/convert-docx",
        json={
            "repo_root": str(tmp_path),
            "source_path": str(src),
            "output_filename": "out.md",
        },
    )
    assert resp.status_code == 400
    assert "docx" in resp.json()["detail"].lower()


def test_convert_docx_endpoint_rejects_missing_source(tmp_path: Path, client) -> None:
    resp = client.post(
        "/api/workspace/convert-docx",
        json={
            "repo_root": str(tmp_path),
            "source_path": str(tmp_path / "nope.docx"),
            "output_filename": "out.md",
        },
    )
    assert resp.status_code == 404


def test_convert_docx_endpoint_rejects_output_traversal(tmp_path: Path, client) -> None:
    src = tmp_path / "resume.docx"
    _build_sample_docx(src)

    resp = client.post(
        "/api/workspace/convert-docx",
        json={
            "repo_root": str(tmp_path),
            "source_path": str(src),
            "output_filename": "../escape.md",
        },
    )
    assert resp.status_code == 400
    detail = resp.json()["detail"].lower()
    assert "escape" in detail or "output_filename" in detail


def test_convert_docx_endpoint_creates_nested_output_dir(tmp_path: Path, client) -> None:
    """Output paths like ``applications/{slug}/resume.md`` auto-create parents."""
    src = tmp_path / "master.docx"
    _build_sample_docx(src)

    resp = client.post(
        "/api/workspace/convert-docx",
        json={
            "repo_root": str(tmp_path),
            "source_path": str(src),
            "output_filename": "applications/acme_corp_senior_engineer/resume.md",
        },
    )
    assert resp.status_code == 200
    out = Path(resp.json()["output_path"])
    assert out.exists()
    assert out.parent.name == "acme_corp_senior_engineer"
    assert "# Jane Q Candidate" in out.read_text(encoding="utf-8")
